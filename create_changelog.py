from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text):
    doc.add_paragraph(text)

def b(text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        para.add_run(text)
    else:
        para.add_run(text)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Light List Accent 1')
    for i, hdr in enumerate(headers):
        t.rows[0].cells[i].text = hdr
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = val

# Title
title = doc.add_heading('FinPal v2.0 — Version 2 Changelog', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p('This document details all changes made in Version 2 of the FinPal Credit Intelligence Platform, based on the enterprise documentation feedback.')

# Meta
tbl(['Field', 'Value'], [
    ['Version', '2.0'],
    ['Previous Version', '1.0'],
    ['Date', 'March 2026'],
    ['Status', 'Released'],
    ['Repository', 'https://github.com/Vinayakcsnci/Finpals'],
])

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────
h('1. Regulatory Positioning & Governance (Gap #1)')
p('The platform now leads with compliance as an architectural pillar, not a bolt-on feature.')
b('Added EU AI Act high-risk classification badge on every application ("EU AI Act: High-Risk")')
b('Added EU AI Act Article references (Art.9-15) to all compliance checks')
b('Added EBA Guidelines on Loan Origination & Monitoring (EBA/GL/2020/06) alignment throughout')
b('Updated topbar subtitle to show "EU AI Act High-Risk Compliant / EBA LOM Aligned"')
b('Added DORA (operational resilience) indicator in compliance dashboard')
b('Added Basel credit risk management principles reference')
b('Updated sidebar footer with jurisdiction indicator, regulation badges (EU AI Act, EBA LOM), and GDPR Art.22 compliance status')

# ─────────────────────────────────────────────────────────────
h('2. GDPR & Automated Decisioning Controls (Gap #2)')
p('Full GDPR Article 22 compliance following the CJEU SCHUFA rulings.')
b('Added GDPR Consent Ledger in the Compliance Checks tab showing all consent records with timestamps')
b('Added consent types: Open Banking (PSD2), Credit Bureau Pull, Automated Processing (GDPR Art.22), Data Retention')
b('Added "GDPR Art.13/14 Notice" as first compliance check for every application')
b('Added "Right to Explanation" button (GDPR Art.22)')
b('Added "Right to Object" button (GDPR Art.21)')
b('Added "Data Subject Request" button')
b('Updated decision modal text to reference GDPR Art.22 requirements')
b('Added consent capture to new application flow')
b('Added SCHUFA ruling compliance references in compliance dashboard')

# ─────────────────────────────────────────────────────────────
h('3. Credit Risk Framework — PD/LGD/EAD (Gap #3)')
p('Replaced the DSCR-only model with a full credit risk framework.')
b('Added PD (Probability of Default) — application model v3.2')
b('Added LGD (Loss Given Default) — model v2.1')
b('Added EAD (Exposure at Default) — model v1.8')
b('Added Affordability Score (capacity-to-repay, cash-flow-based) — model v2.0')
b('Added Risk Grade (A1-D2 internal rating scale)')
b('Added Risk-Based Pricing (APR calculated per risk profile)')
b('Updated Overview tab: now shows 5-card grid (PD, Risk Grade, DSCR, APR, Affordability) + 4-card grid (Revenue, Lodgements, Cash Flow, ECL)')
b('All credit risk data added to application data model for all 4 applications')

# ─────────────────────────────────────────────────────────────
h('4. Model Risk Management / MRM (Gap #4)')
p('Added full Model Governance section accessible from sidebar.')
b('New sidebar nav item: "Model Governance" under Risk & Compliance section')
b('Model Registry showing 6 active models: PD App (v3.2), PD Behavioral (v2.8), LGD (v2.1), EAD (v1.8), Affordability (v2.0), Risk-Based Pricing (v1.5)')
b('Each model shows: version, validation status, last validated date, AUC/Gini, challenger model')
b('SR 11-7 alignment badge displayed')
b('Model Monitoring section with drift indicators: PSI, CSI, Calibration ratio')
b('Framework documentation reference (SR 11-7 / Federal Reserve)')

# ─────────────────────────────────────────────────────────────
h('5. Fairness & Explainability at Scale (Gap #5)')
p('Concrete SHAP/LIME implementation replacing the generic "explainable AI" claim.')
b('New "Risk & XAI" tab for every application')
b('SHAP-based reason codes with visual bar chart (top 5 factors per decision)')
b('Each SHAP factor shows: feature name, direction (positive/negative), magnitude')
b('Color-coded bars: teal for positive factors, red for negative')
b('Fairness metrics displayed: Disparate Impact Ratio and Equal Opportunity Delta')
b('Thresholds shown (DI > 0.80, EO < 0.10) with pass/fail indicators')
b('Model version attribution on all explanations')
b('"GPU-accelerated" label indicating portfolio-scale computation capability')
b('Updated narrative box label from "AI SUMMARY" to "XAI-GENERATED SUMMARY (SHAP v3.2)"')
b('Added "Fairness & Bias Check" as a compliance check item for each application')

# ─────────────────────────────────────────────────────────────
h('6. Credit Rating & Bureau Reporting (Gap #6)')
p('Added credit bureau data furnishing standards and dispute management.')
b('New sidebar nav item: "Bureau Reporting" under Operations section')
b('Bureau furnishing status dashboard (current cycle, error rate, SLA compliance)')
b('Jurisdictional reporting packs: Ireland/EU (ICB), UK (CAIS), US (Metro 2/CRRG), Pan-African')
b('Dispute management queue with open disputes count')
b('Internal rating system panel: A1-D2 scale, migration tracking, override governance')
b('Regulatory reporting: AI Act/HRA documentation, ECL disclosures (IFRS 7/9)')
b('Error-correction SLA tracking (below 1% target)')

# ─────────────────────────────────────────────────────────────
h('7. Collections & Borrower Protection (Gap #7)')
p('Added responsible digital credit safeguards per CGAP guidelines.')
b('New sidebar nav item: "Collections" under Operations section')
b('Responsible Digital Credit Framework banner with CGAP reference')
b('Active collections, hardship cases, and over-indebtedness alerts dashboard')
b('Borrower protection controls: over-indebtedness screening, hardship/rescheduling, contact channel optimization, behavioral nudges with guardrails, consumer protection checks, data misuse prevention')
b('Over-indebtedness warning system (flagged for Ferris Hospitality)')
b('Added "Over-Indebtedness Warning" as compliance check for high-risk applications')

# ─────────────────────────────────────────────────────────────
h('8. Pricing Governance & Collateral Valuation (Gap #8)')
p('Risk-based pricing transparency aligned with EBA LOM expectations.')
b('Added pricing transparency section in Evidence Pack tab')
b('Shows: Risk-Based APR, Risk Grade, PD component spread, LGD capital charge')
b('EBA LOM pricing framework reference')
b('Collateral valuation indicator (independent valuation for Asset Finance; N/A for unsecured)')
b('APR displayed in application overview and evidence header')

# ─────────────────────────────────────────────────────────────
h('9. IFRS 9 / ECL Provisioning (Gap #9 — from Credit Risk Framework)')
p('Full IFRS 9 Expected Credit Loss implementation.')
b('New "IFRS 9" tab for every application')
b('ECL Stage badge (Stage 1/2/3) with color coding')
b('12-month ECL and Lifetime ECL amounts')
b('Scenario overlays: Base (60%), Upside (20%), Downside (20%)')
b('PD/LGD/EAD inputs displayed as point-in-time values')
b('Staging trigger explanation (SICR detection)')
b('Discount rate (EIR) shown')
b('IFRS 9 compliance note referencing EBA LOM expectations and IFRS 7/9 disclosure requirements')
b('Portfolio-level IFRS 9 staging distribution in Portfolio Risk dashboard')
b('ECL (12-month) card added to application Overview tab')

# ─────────────────────────────────────────────────────────────
h('10. Operations Resilience & Country Expansion (Gaps #9 & #10)')
p('DORA alignment and jurisdictional regulatory mapping.')
b('DORA (Digital Operational Resilience Act) compliance indicator in Compliance Dashboard')
b('Jurisdiction selector/indicator in sidebar footer')
b('Regulatory mapping status per jurisdiction in Bureau Reporting')
b('Country-specific configuration: Ireland/EU, UK, US, Pan-African')
b('Data residency indicator maintained (Azure West EU)')

# ─────────────────────────────────────────────────────────────
h('Additional Platform Changes')

h('New Sidebar Navigation', level=2)
b('Risk & Compliance section: Model Governance, Portfolio Risk, Compliance Dashboard')
b('Operations section: Collections, Bureau Reporting')
b('Platform section: Applications (existing), Audit Trail (existing)')

h('New Application Tabs', level=2)
p('Application tabs expanded from 5 to 7:')
b('Overview (enhanced with PD/LGD/EAD/APR/Affordability/ECL)')
b('Risk & XAI (NEW — SHAP explanations, fairness metrics, credit risk detail)')
b('Open Banking (enhanced with affordability score)')
b('IFRS 9 (NEW — ECL staging, scenarios, PD/LGD/EAD inputs)')
b('Compliance (enhanced with GDPR consent ledger, EU AI Act checks, fairness checks, borrower protection)')
b('Evidence Pack (enhanced with pricing transparency, risk grade, APR)')
b('Audit Log (enhanced with GDPR consent entries, model execution, SHAP generation, IFRS 9, EU AI Act verification)')

h('Compliance Dashboard (NEW)', level=2)
b('EU AI Act (High-Risk) compliance matrix: Art.9-15 with status indicators')
b('GDPR compliance matrix: Art.6, 13/14, 22, 25, 35, DSR, SCHUFA ruling')
b('EBA LOM compliance: Credit risk policy, pricing, collateral, monitoring, data infrastructure')
b('Additional frameworks: DORA, MRM/SR 11-7, IFRS 9, Basel, CBI IAF')

h('Portfolio Risk Dashboard (NEW)', level=2)
b('Total exposure, weighted average PD, PAR30/NPL, ECL coverage')
b('IFRS 9 staging distribution (Stage 1/2/3 percentages)')
b('Sector concentration risk analysis')

h('Processing Animation Updated', level=2)
p('Processing steps expanded from 14 to 20 to reflect new capabilities:')
b('Added: Computing PD/LGD/EAD models')
b('Added: Running affordability & capacity-to-repay')
b('Added: Calculating risk-based pricing')
b('Added: Generating SHAP explanations')
b('Added: Evaluating fairness metrics')
b('Added: Computing IFRS 9 ECL staging')
b('Added: Applying EBA LOM credit policy (was generic "lender credit policy")')
b('Added: Validating EU AI Act compliance')

h('Data Model Enhancements', level=2)
p('Each application object now includes:')
b('pd, lgd, ead — Credit risk model outputs')
b('affordability — Capacity-to-repay score (0-1)')
b('riskGrade — Internal rating (A1-D2)')
b('apr — Risk-based annual percentage rate')
b('eclStage, ecl12m, eclLifetime — IFRS 9 ECL data')
b('shapValues[] — Top 5 SHAP factors with feature, value, direction')
b('fairnessMetrics — Disparate impact ratio and equal opportunity delta')
b('consentLog[] — GDPR consent records with type, timestamp, status')

h('Audit Trail Enhanced', level=2)
b('Added: GDPR Art.13/14 notice presented')
b('Added: Consent obtained (GDPR Art.22)')
b('Added: PD/LGD/EAD models executed (with model versions)')
b('Added: SHAP explanations generated (EU AI Act Art.13)')
b('Added: Fairness metrics computed')
b('Added: IFRS 9 ECL calculated')
b('Added: EU AI Act compliance verified')
b('Added: Risk-based pricing calculated (EBA LOM)')

# ─────────────────────────────────────────────────────────────
h('Summary of Changes by Documentation Gap')
tbl(
    ['Gap #', 'Issue', 'Status', 'Changes Made'],
    [
        ['1', 'Regulatory positioning (EU AI Act, EBA LOM)', 'Resolved', 'High-risk badges, Art.9-15 checks, EBA LOM policy gates, DORA'],
        ['2', 'GDPR & automated decisioning', 'Resolved', 'Consent ledger, Art.22 safeguards, right to explanation/object, DSR buttons'],
        ['3', 'Credit risk framework depth', 'Resolved', 'PD/LGD/EAD models, affordability, risk grades, risk-based pricing'],
        ['4', 'Model risk management', 'Resolved', 'Model registry, 6 models, SR 11-7, drift monitoring, challenger models'],
        ['5', 'Fairness & explainability', 'Resolved', 'SHAP bars, fairness metrics (DI/EO), bias checks, model attribution'],
        ['6', 'Credit rating & reporting', 'Resolved', 'Bureau reporting, Metro 2/CRRG, jurisdiction packs, dispute mgmt, internal ratings'],
        ['7', 'Collections & borrower protection', 'Resolved', 'Collections dashboard, CGAP guardrails, over-indebtedness alerts, hardship'],
        ['8', 'Pricing governance & collateral', 'Resolved', 'EBA LOM pricing transparency, APR breakdown, collateral valuation'],
        ['9', 'Operations resilience & third-party risk', 'Resolved', 'DORA compliance, Basel principles, operational controls'],
        ['10', 'Country expansion readiness', 'Resolved', 'Jurisdiction config, regulatory mapping, multi-market bureau packs'],
    ]
)

out = r'C:\Users\lenovo\Downloads\NCI\CITI\feedback\FinPal_V2_Changelog.docx'
doc.save(out)
print(f'Changelog saved to: {out}')
