#!/usr/bin/env python3
"""Generate FinPal_V4_Changelog.docx — Changes from v3.0 → v4.0"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "FinPal_V4_Changelog.docx")

# ── Colour palette ─────────────────────────────────────────────
NAVY    = RGBColor(0x0B, 0x1F, 0x3A)
TEAL    = RGBColor(0x00, 0xA8, 0x96)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
MUTED   = RGBColor(0x7A, 0x94, 0xAC)
GOLD    = RGBColor(0xF5, 0xA6, 0x23)
GREEN   = RGBColor(0x2E, 0xCC, 0x71)
RED     = RGBColor(0xE8, 0x4C, 0x4C)
LIGHT   = RGBColor(0xF0, 0xF6, 0xFA)
DARK_BG = RGBColor(0x0C, 0x22, 0x40)

# ── Helpers ────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    run.font.color.rgb = color
    return p

def body(doc, text, size=10.5, color=None, bold=False, italic=False,
         space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    add_run(p, text, bold=bold, italic=italic, size=size,
            color=color or RGBColor(0x1A, 0x1A, 0x2E))
    return p

def bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, size=size, color=RGBColor(0x1A, 0x1A, 0x2E))
    return p

def divider(doc, color='00A896'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_header_row(table, cols, bg='0B1F3A'):
    row = table.rows[0]
    for i, (text, width) in enumerate(cols):
        cell = row.cells[i]
        cell.width = Inches(width)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, bold=True, size=9, color=WHITE, font='Calibri')

def banner_table(doc, title, subtitle, bg='0B1F3A', accent='00A896'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    add_run(p, title + '\n', bold=True, size=22, color=WHITE, font='Calibri')
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    add_run(p2, subtitle, bold=False, size=10, color=RGBColor(0x00, 0xA8, 0x96), font='Calibri')
    return tbl

# ── Build document ─────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.4)
    section.right_margin  = Cm(2.4)

# ── Cover banner ──────────────────────────────────────────────
banner_table(doc,
    'FinPal v4.0 — Changelog',
    'AI-Powered Loan Intake & Assessment Platform  ·  Changes from v3.0 → v4.0')
doc.add_paragraph()

# ── Document meta ─────────────────────────────────────────────
meta = doc.add_table(rows=1, cols=4)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
meta.style = 'Table Grid'
labels = [
    ('Document', 'FinPal_V4_Changelog'),
    ('Version', 'v4.0'),
    ('Date', '12 March 2026'),
    ('Status', 'Released'),
]
for i, (k, v) in enumerate(labels):
    cell = meta.cell(0, i)
    set_cell_bg(cell, 'F0F6FA')
    p = cell.paragraphs[0]
    add_run(p, k + '\n', bold=True, size=8.5, color=MUTED)
    add_run(p, v, bold=False, size=10, color=NAVY)
doc.add_paragraph()
divider(doc)

# ── Executive Summary ─────────────────────────────────────────
heading(doc, '1. Executive Summary', level=2, color=NAVY)
body(doc,
    'FinPal v4.0 is a feature-completeness update ensuring full coverage of all required product '
    'capabilities across four pillars: Core Product, Compliance & Regulatory, Geographic Coverage, '
    'and Integrations & Tech. This version closes all gaps identified against the FinPal product '
    'feature specification and adds new dedicated UI sections for Integrations, Geographic Coverage, '
    'EU AI Act Technical Documentation, and Vendor Due Diligence.',
    space_after=6)

themes = [
    ('Theme 1', 'Irish Open Banking — Full Bank Trio',
     'Bank of Ireland (BOI) and Permanent TSB (PTSB) added alongside AIB. '
     'All three major Irish retail banks now explicitly covered in the Open Banking tab.'),
    ('Theme 2', 'Compliance Completeness — SEAR, FCA, Technical Docs, VDD',
     'SEAR added alongside IAF. FCA/Consumer Duty explicitly listed. '
     'EU AI Act Technical Documentation (Art.11) section added. '
     'Pre-assembled Vendor Due Diligence Pack section added.'),
    ('Theme 3', 'Geographic Coverage — Dedicated Section',
     'New Ireland / UK / EU coverage cards in Architecture view. '
     'Each market shows its specific regulatory, banking and data residency context.'),
    ('Theme 4', 'Integrations & API Layer — New Section',
     'LOS API, REST API, Webhooks, Data Warehouse Export and BI Connector '
     'all added to the Architecture view with full technical detail.'),
]

tbl = doc.add_table(rows=len(themes) + 1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl, [('Theme', 0.8), ('Title', 2.1), ('Summary', 3.6)], bg='0B1F3A')
for i, (num, title, summary) in enumerate(themes):
    row = tbl.rows[i + 1]
    set_cell_bg(row.cells[0], 'E8F8F5')
    add_run(row.cells[0].paragraphs[0], num, bold=True, size=9, color=TEAL)
    add_run(row.cells[1].paragraphs[0], title, bold=True, size=9, color=NAVY)
    add_run(row.cells[2].paragraphs[0], summary, size=9, color=RGBColor(0x1A, 0x1A, 0x2E))
doc.add_paragraph()
divider(doc)

# ── Theme 1: Irish Open Banking ───────────────────────────────
heading(doc, '2. Theme 1 — Irish Open Banking: Full Bank Trio (AIB / BOI / PTSB)', level=2, color=TEAL)
heading(doc, '2.1 Gap Identified', level=3, color=NAVY)
body(doc,
    'The product specification requires explicit Irish Bank Open Banking coverage for AIB, '
    'Bank of Ireland (BOI), and Permanent TSB (PTSB). Prior to v4.0, only AIB was referenced '
    'in the Open Banking tab, leaving the two other major Irish retail banks absent from the UI.')

heading(doc, '2.2 Changes Implemented in v4.0', level=3, color=NAVY)
body(doc, 'A. Open Banking Tab — New Irish Bank Coverage Banner', bold=True, size=10)
for item in [
    'New informational flag box added at the top of the Open Banking tab: '
    '"Irish Open Banking Coverage — AIB · Bank of Ireland · PTSB"',
    'Explains that all three banks are connected via PSD2/AISP authorisation',
    'Clarifies that multi-bank aggregation is supported for borrowers with accounts across institutions',
]:
    bullet(doc, item)

body(doc, 'B. Three-Card Bank Coverage Grid (new in Open Banking tab)', bold=True, size=10, space_before=8)
banks = [
    ('AIB (Allied Irish Banks)', 'PSD2/AISP Connected',
     'Business Current · Savings · Overdraft · Instant pull · 24-month history'),
    ('Bank of Ireland (BOI)', 'BOI Open Finance Certified',
     'Current · Deposit · Loan accounts · 24-month history · Instant pull'),
    ('Permanent TSB (PTSB)', 'PSD2/AISP Connected',
     'Current · Business accounts · 24-month history · On-demand refresh'),
]
tbl2 = doc.add_table(rows=len(banks) + 1, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl2, [('Bank', 1.5), ('Status', 2.0), ('Accounts Covered', 3.0)], bg='0B1F3A')
for i, (bank, status, detail) in enumerate(banks):
    row = tbl2.rows[i + 1]
    set_cell_bg(row.cells[0], 'E8F8F5')
    add_run(row.cells[0].paragraphs[0], bank, bold=True, size=9, color=TEAL)
    add_run(row.cells[1].paragraphs[0], status, bold=True, size=9, color=GREEN)
    add_run(row.cells[2].paragraphs[0], detail, size=9)

body(doc, 'C. Architecture View — Integrations table now lists all three Irish banks explicitly.',
     italic=True, size=9.5, space_before=8)
doc.add_paragraph()
divider(doc)

# ── Theme 2: Compliance Completeness ─────────────────────────
heading(doc, '3. Theme 2 — Compliance Completeness', level=2, color=TEAL)
heading(doc, '3.1 Gap Identified', level=3, color=NAVY)
body(doc,
    'Four compliance gaps were identified against the v4.0 feature specification:')
for item in [
    'SEAR (Senior Executive Accountability Regime) not explicitly named — only "IAF" appeared',
    'FCA (UK) compliance not listed as a standalone regulatory framework',
    'EU AI Act Technical Documentation (Art.11) had no dedicated UI section despite being a '
    'mandatory requirement for high-risk AI systems',
    'Pre-assembled Vendor Due Diligence Pack — a key sales and procurement enabler — was absent',
]:
    bullet(doc, item)

heading(doc, '3.2 Change A — SEAR Added Alongside IAF', level=3, color=NAVY)
body(doc,
    'In the Compliance Dashboard "Additional Frameworks" card, the entry "CBI IAF (Ireland)" '
    'was updated to "CBI IAF / SEAR (Ireland)" with the detail: '
    '"Named PCF holder logged · decision audit trails active". '
    'SEAR (Senior Executive Accountability Regime) is the Irish equivalent of the UK SM&CR '
    'and requires named Pre-Approval Controlled Function holders to be recorded against '
    'every significant credit decision — now explicitly surfaced in the UI.')

heading(doc, '3.3 Change B — FCA (UK) / Consumer Duty Explicitly Listed', level=3, color=NAVY)
body(doc,
    'A new compliance row was added to the "Additional Frameworks" card: '
    '"FCA (UK) — CGAP / Consumer Duty" with status "Responsible lending guardrails · Consumer Duty aligned". '
    'This makes UK regulatory coverage explicit and aligns with the FCA\'s Consumer Duty (effective July 2023) '
    'which requires firms to deliver good outcomes for retail customers.')

heading(doc, '3.4 Change C — EU AI Act Technical Documentation (Art.11) Section', level=3, color=NAVY)
body(doc,
    'A new dedicated section "EU AI Act Technical Documentation (Art.11)" was added to the '
    'Compliance Dashboard. Article 11 of the EU AI Act 2024/1689 requires providers of high-risk '
    'AI systems to maintain detailed technical documentation before placing the system on the market. '
    'The new section documents six components:')

art11_items = [
    ('System Description & Purpose',
     'High-risk AI system — credit scoring & automated underwriting · Annex IV compliant'),
    ('Training Data & Validation',
     'Feature store lineage · PII anonymisation · data quality controls documented'),
    ('Model Architecture & Logic',
     'SHAP-explained PD/LGD/EAD models · challenger framework documented'),
    ('Performance Monitoring Plan',
     'PSI/CSI drift metrics · monthly calibration reports · scheduled revalidation'),
    ('Human Oversight Procedures',
     'Mandatory underwriter gate · escalation paths · override logging documented'),
    ('Post-Market Monitoring',
     'Automated alerts on model drift · adverse outcome reporting pipeline'),
]
tbl3 = doc.add_table(rows=len(art11_items) + 1, cols=2)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl3, [('Art.11 Component', 2.3), ('FinPal Implementation', 4.2)], bg='0B1F3A')
for i, (comp, impl) in enumerate(art11_items):
    row = tbl3.rows[i + 1]
    set_cell_bg(row.cells[0], 'E8F8F5')
    add_run(row.cells[0].paragraphs[0], comp, bold=True, size=9, color=NAVY)
    add_run(row.cells[1].paragraphs[0], impl, size=9, color=GREEN)

heading(doc, '3.5 Change D — Pre-Assembled Vendor Due Diligence Pack', level=3, color=NAVY)
body(doc,
    'A new section "Pre-Assembled Vendor Due Diligence Pack" was added to the Compliance Dashboard '
    'with an informational flag box explaining that the pack is available for lender and regulator review. '
    'The pack covers seven document categories:')

vdd_items = [
    ('Information Security Policy', 'ISO 27001-aligned · last updated Mar 2025'),
    ('GDPR Data Processing Agreement (DPA)', 'Standard contractual clauses · Art.28 processor terms · DPIA summary'),
    ('EU AI Act Technical File (Annex IV)', 'High-risk AI system documentation · CE marking readiness'),
    ('Business Continuity & DR Plan', 'RPO: 4h · RTO: 8h · annual test completed'),
    ('Penetration Test Report (latest)', 'External pen test Q4 2024 · all critical findings remediated'),
    ('Sub-processor Register', 'Azure West EU (primary) · no data transferred outside EEA'),
    ('SOC 2 Type II (in progress)', 'Target certification Q3 2025 · bridging controls in place'),
]
tbl4 = doc.add_table(rows=len(vdd_items) + 1, cols=2)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl4, [('Document', 2.5), ('Status / Detail', 4.0)], bg='0B1F3A')
for i, (doc_name, detail) in enumerate(vdd_items):
    row = tbl4.rows[i + 1]
    set_cell_bg(row.cells[0], 'E8F8F5' if i < 6 else 'FFF8E7')
    add_run(row.cells[0].paragraphs[0], doc_name, bold=True, size=9, color=NAVY)
    c = GREEN if i < 6 else GOLD
    add_run(row.cells[1].paragraphs[0], detail, size=9, color=c)
doc.add_paragraph()
divider(doc)

# ── Theme 3: Geographic Coverage ─────────────────────────────
heading(doc, '4. Theme 3 — Geographic Coverage: Dedicated Section', level=2, color=TEAL)
heading(doc, '4.1 Gap Identified', level=3, color=NAVY)
body(doc,
    'The product specification requires explicit coverage for three geographic markets: Ireland, '
    'United Kingdom, and EU / Europe. While references to these markets existed throughout the UI, '
    'there was no dedicated section that consolidated the regulatory, banking, and data residency '
    'context for each market in one place.')

heading(doc, '4.2 Changes Implemented in v4.0', level=3, color=NAVY)
body(doc,
    'A new "Geographic Coverage" three-card section was added to the Architecture & Security view, '
    'appearing before the Integrations & API Layer. Each card summarises the market-specific stack:')

geo_items = [
    ('Ireland (Primary Market)',
     'CBI-supervised lenders · CRO integration · ROS tax clearance · AIB/BOI/PTSB open banking · '
     'ICB credit bureau · IAF/SEAR compliance · Azure West EU data residency'),
    ('United Kingdom (Active)',
     'FCA-regulated lenders · Consumer Duty aligned · Open Banking (FCA/CMA9) · '
     'Experian/Equifax/TransUnion · CAIS bureau format · UK data residency option'),
    ('EU / Europe (Available)',
     'EU AI Act 2024/1689 native · GDPR-native architecture · EBA LOM aligned · '
     'PSD2 open banking · DORA compliant · EEA data residency · Local bureau connectors on request'),
]
tbl5 = doc.add_table(rows=len(geo_items) + 1, cols=2)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl5, [('Market', 2.0), ('Coverage Detail', 4.5)], bg='0B1F3A')
for i, (market, detail) in enumerate(geo_items):
    row = tbl5.rows[i + 1]
    set_cell_bg(row.cells[0], 'E8F8F5')
    add_run(row.cells[0].paragraphs[0], market, bold=True, size=9, color=TEAL)
    add_run(row.cells[1].paragraphs[0], detail, size=9)
doc.add_paragraph()
divider(doc)

# ── Theme 4: Integrations & API Layer ────────────────────────
heading(doc, '5. Theme 4 — Integrations & API Layer: New Section', level=2, color=TEAL)
heading(doc, '5.1 Gap Identified', level=3, color=NAVY)
body(doc,
    'The product specification requires the following integration and technology capabilities to be '
    'explicitly visible in the product: LOS (Loan Origination System) API, REST API / Webhooks, '
    'Irish Bank Open Banking (AIB/BOI/PTSB), CRO, ROS, Credit Bureau Integration, '
    'Data Warehouse / BI Export, and Private Deployment Option. '
    'Prior to v4.0, these were implicit in the codebase but not surfaced in a dedicated UI section.')

heading(doc, '5.2 Changes Implemented in v4.0', level=3, color=NAVY)
body(doc,
    'A new "Integrations & API Layer" two-column section was added to the Architecture view. '
    'The left card covers core integrations; the right card covers API & data export capabilities:')

core_integrations = [
    ('LOS API (Loan Origination System)', 'REST · bidirectional sync · status webhooks'),
    ('Irish Open Banking — AIB', 'PSD2/AISP · 24-month transaction history'),
    ('Irish Open Banking — Bank of Ireland', 'PSD2/AISP · BOI Open Finance certified'),
    ('Irish Open Banking — PTSB', 'PSD2/AISP · Permanent TSB Open Banking'),
    ('CRO (Companies Registration Office)', 'Company search · director registry · filings'),
    ('ROS (Revenue Online Service)', 'Tax clearance verify · VAT/PAYE/CT status'),
    ('Credit Bureau — ICB (Ireland)', 'Pull & furnish · ICB local schema'),
    ('Credit Bureau — Experian/Equifax/TransUnion', 'UK CAIS format · configurable'),
]

api_capabilities = [
    ('REST API', 'Full CRUD · OpenAPI 3.1 spec · sandbox available'),
    ('Webhooks', 'Event-driven · decision, status, document events'),
    ('Data Warehouse Export', 'Nightly delta exports · Parquet/CSV · S3/Azure Blob'),
    ('BI Connector', 'Power BI & Tableau certified connectors'),
    ('Evidence Pack Export', 'Structured PDF · machine-readable JSON alongside'),
    ('SFTP / Batch', 'Scheduled bulk file drops for legacy LOS integrations'),
    ('Private Deployment', 'On-premise or dedicated cloud · full API parity'),
    ('API Authentication', 'OAuth 2.0 / mTLS · per-tenant credentials'),
]

body(doc, 'Core Integrations Card:', bold=True, size=10, space_before=6)
tbl6 = doc.add_table(rows=len(core_integrations) + 1, cols=2)
tbl6.style = 'Table Grid'
tbl6.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl6, [('Integration', 2.5), ('Detail', 4.0)], bg='0B1F3A')
for i, (intg, detail) in enumerate(core_integrations):
    row = tbl6.rows[i + 1]
    set_cell_bg(row.cells[0], 'E8F8F5')
    add_run(row.cells[0].paragraphs[0], intg, bold=True, size=9, color=NAVY)
    add_run(row.cells[1].paragraphs[0], detail, size=9)

body(doc, 'API & Data Export Card:', bold=True, size=10, space_before=10)
tbl7 = doc.add_table(rows=len(api_capabilities) + 1, cols=2)
tbl7.style = 'Table Grid'
tbl7.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl7, [('Capability', 2.5), ('Detail', 4.0)], bg='0B1F3A')
for i, (cap, detail) in enumerate(api_capabilities):
    row = tbl7.rows[i + 1]
    set_cell_bg(row.cells[0], 'F0F6FA')
    add_run(row.cells[0].paragraphs[0], cap, bold=True, size=9, color=NAVY)
    add_run(row.cells[1].paragraphs[0], detail, size=9)
doc.add_paragraph()
divider(doc)

# ── File Changes Summary ──────────────────────────────────────
heading(doc, '6. Files Changed (v3.0 → v4.0)', level=2, color=NAVY)

file_changes = [
    ('index.html', 'MODIFIED',
     '+89 lines / -5 lines. All changes implemented directly in index.html: '
     'Open Banking tab (BOI/PTSB cards), Compliance Dashboard (SEAR, FCA, Art.11 docs, VDD pack), '
     'Architecture view (Geographic Coverage section, Integrations & API Layer section). '
     'Title updated from v2.0 to v2.1.'),
    ('generate_v4_changelog_docx.py', 'NEW',
     'This changelog generation script — documents all v3.0 → v4.0 changes.'),
    ('FinPal_V4_Changelog.docx', 'NEW',
     'Output of this script — the compiled changelog document.'),
]

tbl8 = doc.add_table(rows=len(file_changes) + 1, cols=3)
tbl8.style = 'Table Grid'
tbl8.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl8, [('File', 1.5), ('Status', 1.0), ('Changes', 4.0)], bg='0B1F3A')
status_colors = {'NEW': TEAL, 'MODIFIED': GOLD, 'REBUILT': GREEN, 'NO CHANGE': MUTED}
for i, (fname, status, changes) in enumerate(file_changes):
    row = tbl8.rows[i + 1]
    bg = 'E8F8F5' if 'MODIFIED' in status or 'NEW' in status else 'F9F9F9'
    set_cell_bg(row.cells[0], bg)
    add_run(row.cells[0].paragraphs[0], fname, bold=True, size=9, color=NAVY)
    sc = status_colors.get(status, MUTED)
    add_run(row.cells[1].paragraphs[0], status, bold=True, size=9, color=sc)
    add_run(row.cells[2].paragraphs[0], changes, size=9)
doc.add_paragraph()
divider(doc)

# ── Feature Coverage Matrix ───────────────────────────────────
heading(doc, '7. Complete Feature Coverage Matrix (v4.0)', level=2, color=NAVY)
body(doc,
    'The following matrix confirms complete coverage of all required features in FinPal v4.0:',
    space_after=6)

coverage = [
    # (Category, Feature, Status, Notes)
    # Core Product
    ('Core Product', 'AI-Powered Underwriting Automation', 'Complete', 'PD/LGD/EAD models, SHAP, HITL decision gate'),
    ('Core Product', 'Document Ingestion & OCR', 'Complete', 'IDP pipeline, OCR extraction, NER, authenticity signals'),
    ('Core Product', 'Open Banking Integration', 'Complete', 'AIB / BOI / PTSB — PSD2/AISP, 24-month history'),
    ('Core Product', 'Transaction Categorisation', 'Complete', 'Revenue, Payroll, Debt Service, Tax categories'),
    ('Core Product', 'DSCR / Cash Flow Calculation', 'Complete', 'DSCR policy minimum 1.20x, free cash flow p.a.'),
    ('Core Product', 'Revenue Reconciliation (declared vs actual)', 'Complete', 'Open banking lodgements vs declared, >5% flag'),
    ('Core Product', 'Automated Credit Scoring / Decisioning', 'Complete', 'A1–C2 risk grades, risk-based pricing (EBA LOM)'),
    ('Core Product', 'Decision Support (human-in-the-loop)', 'Complete', 'Mandatory underwriter gate, rationale capture'),
    ('Core Product', 'Structured Evidence Pack Output', 'Complete', 'PDF + JSON export, all source data included'),
    # Compliance
    ('Compliance', 'GDPR-Native Architecture (EU data residency)', 'Complete', 'Art.6/13/22/25/35 controls, PII Vault, DPA'),
    ('Compliance', 'EU AI Act 2024/1689 Compliance', 'Complete', 'Art.9–15 controls, high-risk classification'),
    ('Compliance', 'CBI (Central Bank of Ireland) Awareness', 'Complete', 'CBI notified, inspection-ready audit logs'),
    ('Compliance', 'FCA (UK) Compliance', 'Complete', 'Consumer Duty aligned, CGAP responsible lending'),
    ('Compliance', 'AML / PEP Screening', 'Complete', 'All directors screened at origination'),
    ('Compliance', 'Immutable Audit Trail', 'Complete', 'Cryptographically immutable, timestamped, Art.12'),
    ('Compliance', 'IAF / SEAR Personal Liability Protection', 'Complete', 'Named PCF holder logged, decision audit trail'),
    ('Compliance', 'Pre-assembled Vendor Due Diligence Pack', 'Complete', '7 documents: InfoSec, DPA, AI Tech File, BCP, PenTest, Sub-processors, SOC2'),
    ('Compliance', 'EU AI Act Technical Documentation', 'Complete', 'Art.11 Annex IV — 6 components documented in UI'),
    # Geography
    ('Geography', 'Ireland Market', 'Complete', 'CBI, CRO, ROS, AIB/BOI/PTSB, ICB, IAF/SEAR, Azure West EU'),
    ('Geography', 'United Kingdom', 'Complete', 'FCA, Consumer Duty, CMA9 Open Banking, Experian/Equifax/TU'),
    ('Geography', 'EU / Europe', 'Complete', 'EU AI Act, GDPR, EBA LOM, PSD2, DORA, EEA data residency'),
    # Integrations
    ('Integrations', 'LOS (Loan Origination System) API', 'Complete', 'REST, bidirectional sync, status webhooks'),
    ('Integrations', 'REST API / Webhooks', 'Complete', 'OpenAPI 3.1, event-driven webhooks, OAuth 2.0 / mTLS'),
    ('Integrations', 'Irish Bank Open Banking (AIB/BOI/PTSB)', 'Complete', 'All three major Irish banks, PSD2/AISP'),
    ('Integrations', 'CRO (Companies Registration Office)', 'Complete', 'Company search, director registry, filing history'),
    ('Integrations', 'ROS (Revenue Online Service) Verify', 'Complete', 'VAT, PAYE, CT tax clearance verification'),
    ('Integrations', 'Credit Bureau Integration', 'Complete', 'ICB (Ireland), Experian/Equifax/TU (UK)'),
    ('Integrations', 'Data Warehouse / BI Export', 'Complete', 'Parquet/CSV, S3/Azure Blob, Power BI & Tableau connectors'),
    ('Integrations', 'Private Deployment Option', 'Complete', 'On-premise and dedicated cloud, full API parity'),
]

tbl9 = doc.add_table(rows=len(coverage) + 1, cols=4)
tbl9.style = 'Table Grid'
tbl9.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl9, [('Category', 1.1), ('Feature', 2.0), ('Status', 0.85), ('Notes', 2.55)], bg='0B1F3A')
cat_colors = {
    'Core Product': 'E8F8F5',
    'Compliance': 'FFF8E7',
    'Geography': 'EAF4FB',
    'Integrations': 'F4ECF7',
}
for i, (cat, feat, status, notes) in enumerate(coverage):
    row = tbl9.rows[i + 1]
    set_cell_bg(row.cells[0], cat_colors.get(cat, 'F9F9F9'))
    add_run(row.cells[0].paragraphs[0], cat, bold=True, size=8.5, color=NAVY)
    add_run(row.cells[1].paragraphs[0], feat, size=9)
    add_run(row.cells[2].paragraphs[0], status, bold=True, size=9, color=GREEN)
    add_run(row.cells[3].paragraphs[0], notes, size=8.5, color=RGBColor(0x33, 0x4E, 0x68))
doc.add_paragraph()
divider(doc)

# ── What Was Preserved ────────────────────────────────────────
heading(doc, '8. What Was Not Changed (Intentionally Preserved from v3.0)', level=2, color=NAVY)
body(doc,
    'The following v3.0 features were intentionally left unchanged as they remain complete and correct:')
preserved = [
    'All 7 application tabs: Overview, Risk & XAI, Open Banking, IFRS 9, Compliance, Evidence Pack, Audit Log',
    'SHAP explainability engine (EU AI Act Art.13) — full local and global explanations',
    'GDPR consent ledger (Art.13/14/22 safeguards, SCHUFA ruling, data subject rights)',
    'EBA LOM risk-based pricing framework (APR calculation, pricing grid)',
    'IFRS 9 ECL engine (Stage 1/2/3, multi-scenario PD overlays, SICR tracking)',
    'Fairness metrics (Disparate Impact Ratio > 0.80, Equal Opportunity Delta < 0.10)',
    'Immutable audit trail (EU AI Act Art.12, CBI IAF, 22+ processing steps logged)',
    'Model Governance / MRM framework (SR 11-7 alignment, challenger models, PSI/CSI drift)',
    'Portfolio Risk Dashboard (concentration risk, IFRS 9 staging distribution)',
    'Collections & Borrower Protection (CGAP responsible digital credit)',
    'Bureau Reporting (ICB Ireland, CAIS UK, Metro 2 US, Pan-African packs)',
    'Intake Pipeline view (5-step borrower journey, IDP as hero feature)',
    'Architecture & Security view (tenant isolation, private AI, security controls)',
    'All 4 sample applications: O\'Brien Construction, Clancy Logistics, Ferris Hospitality, GreenwayTech',
    'Tenant-isolated SaaS architecture with on-premise deployment option',
    'Processing animation with 22 IDP pipeline step labels',
]
for item in preserved:
    bullet(doc, item)
doc.add_paragraph()
divider(doc)

# ── Version History ───────────────────────────────────────────
heading(doc, '9. Version History', level=2, color=NAVY)
ver_tbl = doc.add_table(rows=5, cols=3)
ver_tbl.style = 'Table Grid'
add_header_row(ver_tbl, [('Version', 0.8), ('Date', 1.5), ('Summary', 4.7)], bg='0B1F3A')
versions = [
    ('v1.0', 'Dec 2024',
     'Initial upload — basic loan application prototype'),
    ('v2.0', 'Feb 2025',
     'Full platform upgrade: EU AI Act, EBA LOM, GDPR, IFRS 9 compliance overhaul'),
    ('v3.0', 'Mar 2026',
     'Product reframe: IDP pipeline, tenant isolation, private AI architecture, SaaS data governance'),
    ('v4.0', '12 Mar 2026',
     'Feature completeness: BOI/PTSB open banking, SEAR, FCA, EU AI Act Art.11 Tech Docs, '
     'Vendor Due Diligence Pack, Geographic Coverage section, LOS API, REST API, Webhooks, '
     'Data Warehouse/BI Export'),
]
for i, (v, d, s) in enumerate(versions):
    row = ver_tbl.rows[i + 1]
    set_cell_bg(row.cells[0], 'E8F8F5' if i == 3 else 'F9F9F9')
    add_run(row.cells[0].paragraphs[0], v, bold=True, size=9, color=TEAL if i == 3 else NAVY)
    add_run(row.cells[1].paragraphs[0], d, size=9)
    add_run(row.cells[2].paragraphs[0], s, size=9)

doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_footer,
    'FinPal v4.0  ·  Feature Completeness Release  ·  12 March 2026  ·  Confidential',
    size=8.5, color=MUTED, italic=True)

# ── Save ──────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT):,} bytes")
